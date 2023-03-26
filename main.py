import Resume
import tkinter as tk
import tkinter.filedialog as filedialog
import tkinter.messagebox as messagebox
from tkinter import font
from docx import Document
from docx.shared import Pt
import tkinter.scrolledtext as tkst

class Application(tk.Frame):
    name = city = state = zip_code = phone = email = job = skills = education = ''
    experiences = []

    def __init__(self, master=None):
        super().__init__(master)
        self.cv = ''
        self.master = master
        self.pack()
        self.create_widgets()
        self.master.geometry(f"750x750")
        self.adjustedResume = ''
        self.font = font.nametofont("TkDefaultFont")
        self.update()
        self.master.title("GPTExpert - AI Resume Tool (resumepros123)")
        self.master.iconbitmap("gptexpert.ico")


    def update_geometry(self):
        # Update the geometry of the window
        self.master.geometry(f"{self.winfo_reqwidth()+100}x{self.winfo_reqheight()+100}")

    def create_widgets(self):
        # create label
        self.label = tk.Label(self, text="Chose Resume Level", font=("Arial", 18), fg="navy")
        self.label.grid(row=0, column=0, padx=10, pady=10, sticky=tk.W)

        # create radio buttons
        self.level = tk.StringVar()
        self.radio1 = tk.Radiobutton(self, text="Entry Level", variable=self.level, value="Entry Level",
                                     font=("Arial", 14), fg="navy")
        self.radio2 = tk.Radiobutton(self, text="Mid Level", variable=self.level, value="Mid Level", font=("Arial", 14),
                                     fg="navy")
        self.radio3 = tk.Radiobutton(self, text="Senior Level", variable=self.level, value="Senior Level",
                                     font=("Arial", 14), fg="navy")
        self.radio4 = tk.Radiobutton(self, text="Executive Level", variable=self.level, value="Executive Level",
                                     font=("Arial", 14), fg="navy")
        self.radio1.grid(row=1, column=0, padx=(30, 10), pady=10, sticky=tk.W)
        self.radio2.grid(row=2, column=0, padx=(30, 10), pady=10, sticky=tk.W)
        self.radio3.grid(row=3, column=0, padx=(30, 10), pady=10, sticky=tk.W)
        self.radio4.grid(row=4, column=0, padx=(30, 10), pady=10, sticky=tk.W)

        self.submit = tk.Button(self, text="Next", command=self.submit_choice, font=("Arial", 14), bg="navy",
                                fg="white")
        self.submit.grid(row=19, column=0, padx=10, pady=20, sticky=tk.N, ipadx=50)

    def submit_choice(self):
        self.resumeLevel = self.level.get()

        # destroy current widgets
        self.label.destroy()
        self.radio1.destroy()
        self.radio2.destroy()
        self.radio3.destroy()
        self.radio4.destroy()
        self.submit.destroy()

        tk.Label(self, text="Fill in client information (if not applicable, leave empty):", font=("Arial", 14), fg="navy").grid(row=0, column=0, padx=10, pady=10, sticky=tk.W)

        # create new widgets for the next form
        tk.Label(self, text="Name:", font=("Arial", 14), fg="navy").grid(row=1, column=0, padx=10, pady=0, sticky=tk.W)
        self.nameB = tk.Text(self, height=1, width=30, font=("Arial", 18), bg="#F9F8F8")
        self.nameB.grid(row=2, column=0, padx=10, pady=0)

        tk.Label(self, text="City:", font=("Arial", 14), fg="navy").grid(row=3, column=0, padx=10, pady=0, sticky=tk.W)
        self.cityB = tk.Text(self, height=1, width=30, font=("Arial", 18), bg="#F9F8F8")
        self.cityB.grid(row=4, column=0, padx=10, pady=0)

        tk.Label(self, text="State:", font=("Arial", 14), fg="navy").grid(row=5, column=0, padx=10, pady=0,
                                                                          sticky=tk.W)
        self.stateB = tk.Text(self, height=1, width=30, font=("Arial", 14), bg="#F9F8F8")
        self.stateB.grid(row=6, column=0, padx=10, pady=0)

        tk.Label(self, text="Zip Code:", font=("Arial", 14), fg="navy").grid(row=7, column=0, padx=10, pady=0,
                                                                             sticky=tk.W)
        self.zip_codeB = tk.Text(self, height=1, width=30, font=("Arial", 14), bg="#F9F8F8")
        self.zip_codeB.grid(row=8, column=0, padx=10, pady=0)

        tk.Label(self, text="Phone:", font=("Arial", 14), fg="navy").grid(row=9, column=0, padx=10, pady=0,
                                                                          sticky=tk.W)
        self.phoneB = tk.Text(self, height=1, width=30, font=("Arial", 14), bg="#F9F8F8")
        self.phoneB.grid(row=10, column=0, padx=10, pady=0)

        tk.Label(self, text="Email:", font=("Arial", 14), fg="navy").grid(row=11, column=0, padx=10, pady=0,
                                                                          sticky=tk.W)
        self.emailB = tk.Text(self, height=1, width=30, font=("Arial", 14), bg="#F9F8F8")
        self.emailB.grid(row=12, column=0, padx=10, pady=0)

        tk.Label(self, text="Job title:", font=("Arial", 14), fg="navy").grid(row=13, column=0, padx=10, pady=0, sticky=tk.W)
        self.jobB = tk.Text(self, height=1, width=30, font=("Arial", 14), bg="#F9F8F8")
        self.jobB.grid(row=14, column=0, padx=10, pady=0)

        tk.Label(self, text="Skills:", font=("Arial", 14), fg="navy").grid(row=15, column=0, padx=10, pady=0,
                                                                           sticky=tk.W)
        self.skillsB = tk.Text(self, height=3, width=30, font=("Arial", 14), bg="#F9F8F8", wrap="word")
        self.skillsB.grid(row=16, column=0, padx=10, pady=0)
        scrollb = tk.Scrollbar(self, command=self.skillsB.yview)
        self.skillsB['yscrollcommand'] = scrollb.set
        scrollb.grid(row=16, column=1, pady=0, sticky=tk.NS)

        tk.Label(self, text="Education:", font=("Arial", 14), fg="navy").grid(row=17, column=0, padx=10, pady=0,
                                                                              sticky=tk.W)
        self.educationB = tk.Text(self, height=3, width=30, font=("Arial", 14), bg="#F9F8F8", wrap="word")
        self.educationB.grid(row=18, column=0, padx=10, pady=0)
        scrollb = tk.Scrollbar(self, command=self.educationB.yview)
        self.educationB['yscrollcommand'] = scrollb.set
        scrollb.grid(row=18, column=1, pady=0, sticky=tk.NS)

        # create submit button for the next form
        self.submit2 = tk.Button(self, text="Next", command=self.submit_choice2, font=("Arial", 14), bg="navy",
                                 fg="white")
        self.submit2.grid(row=19, column=0, padx=10, pady=0, sticky=tk.N, ipadx=50)

    def submit_choice2(self):
        # retrieve data from the fields
        self.name = self.nameB.get("1.0",'end-1c')
        self.city = self.cityB.get("1.0",'end-1c')
        self.state = self.stateB.get("1.0",'end-1c')
        self.zip_code = self.zip_codeB.get("1.0",'end-1c')
        self.phone = self.phoneB.get("1.0",'end-1c')
        self.email = self.emailB.get("1.0",'end-1c')
        self.job = self.jobB.get("1.0",'end-1c')
        self.skills = self.skillsB.get("1.0",'end-1c')
        self.education = self.educationB.get("1.0",'end-1c')
        self.clear_fields()

    def clear_fields(self):
        # destroy widgets
        for widgets in self.winfo_children():
            widgets.destroy()
        self.experience_choice()

    def experience_choice(self):
        # create label
        self.label = tk.Label(self, text="Specify the number of the client's experiences", font=("Arial", 14), fg="navy")
        self.label.grid(row=0, column=0, padx=10, pady=10, sticky=tk.W)

        # create radio buttons
        self.choice = tk.IntVar()
        self.radio1 = tk.Radiobutton(self, text="1 Experience", variable=self.choice, value=1, font=("Arial", 14), fg="navy")
        self.radio2 = tk.Radiobutton(self, text="2 Experiences", variable=self.choice, value=2, font=("Arial", 14), fg="navy")
        self.radio3 = tk.Radiobutton(self, text="3 Experiences", variable=self.choice, value=3, font=("Arial", 14), fg="navy")
        self.radio4 = tk.Radiobutton(self, text="4 Experiences", variable=self.choice, value=4, font=("Arial", 14), fg="navy")
        self.radio1.grid(row=1, column=0, padx=10, pady=5, sticky=tk.W)
        self.radio2.grid(row=2, column=0, padx=10, pady=5, sticky=tk.W)
        self.radio3.grid(row=3, column=0, padx=10, pady=5, sticky=tk.W)
        self.radio4.grid(row=4, column=0, padx=10, pady=5, sticky=tk.W)

        # create submit button
        self.submit3 = tk.Button(self, text="Next", command=self.submit_choice3, font=("Arial", 14), bg="navy",
                                 fg="white")
        self.submit3.grid(row=5, column=0, padx=10, pady=20, sticky=tk.N, ipadx=50)

    def submit_choice3(self):
        self.num = self.choice.get()
        self.label.destroy()
        self.radio1.destroy()
        self.radio2.destroy()
        self.radio3.destroy()
        self.radio4.destroy()
        self.submit3.destroy()
        self.experience()

    def experience(self):
        tk.Label(self, text="Fill in experience information (leave empty if not applicable), press next to go to the "
                            "next experience or press create resume",
                 font=("Arial", 18), fg="navy").grid(row=0, column=0, padx=10, pady=10, sticky=tk.W)

        tk.Label(self, text="Company Name:", font=("Arial", 14), fg="navy").grid(row=1, column=0, padx=10, pady=10, sticky=tk.W)
        self.comp = tk.Text(self, height=1, width=30, font=("Arial", 14), bg="#F9F8F8")
        self.comp.grid(row=2, column=0, padx=10, pady=5)

        tk.Label(self, text="Job Name:", font=("Arial", 14), fg="navy").grid(row=3, column=0, padx=10, pady=10, sticky=tk.W)
        self.job = tk.Text(self, height=1, width=30, font=("Arial", 14), bg="#F9F8F8")
        self.job.grid(row=4, column=0, padx=10, pady=5)

        tk.Label(self, text="Location:", font=("Arial", 14), fg="navy").grid(row=5, column=0, padx=10, pady=10, sticky=tk.W)
        self.location = tk.Text(self, height=1, width=30, font=("Arial", 14), bg="#F9F8F8")
        self.location.grid(row=6, column=0, padx=10, pady=5)

        tk.Label(self, text="Dates:", font=("Arial", 14), fg="navy").grid(row=7, column=0, padx=10, pady=10, sticky=tk.W)
        self.dates = tk.Text(self, height=1, width=30, font=("Arial", 14), bg="#F9F8F8")
        self.dates.grid(row=8, column=0, padx=10, pady=5)

        tk.Label(self, text="Company Description:", font=("Arial", 14), fg="navy").grid(row=9, column=0, padx=10, pady=10,
                                                                             sticky=tk.W)
        self.compDesc = tk.Text(self, height=3, width=30, font=("Arial", 14), bg="#F9F8F8", wrap="word")
        self.compDesc.grid(row=10, column=0, padx=10, pady=5)

        tk.Label(self, text="Job Description:", font=("Arial", 14), fg="navy").grid(row=11, column=0, padx=10, pady=10,
                                                                         sticky=tk.W)
        self.jobDesc = tk.Text(self, height=3, width=30, font=("Arial", 14), bg="#F9F8F8", wrap="word")
        self.jobDesc.grid(row=12, column=0, padx=10, pady=5)
        scrollb1 = tk.Scrollbar(self, command=self.jobDesc.yview)
        self.jobDesc['yscrollcommand'] = scrollb1.set
        scrollb1.grid(row=12, column=1, pady=5, sticky=tk.NS)

        tk.Label(self, text="Achievements:", font=("Arial", 14), fg="navy").grid(row=13, column=0, padx=10, pady=10, sticky=tk.W)
        self.achievements = tk.Text(self, height=3, width=30, font=("Arial", 14), bg="#F9F8F8", wrap="word")
        self.achievements.grid(row=14, column=0, padx=10, pady=5)
        scrollb2 = tk.Scrollbar(self, command=self.achievements.yview)
        self.achievements['yscrollcommand'] = scrollb2.set
        scrollb2.grid(row=14, column=1, pady=5, sticky=tk.NS)

        # create submit buttons
        self.submit4 = tk.Button(self, text="Next", command=self.submit_choice4, font=("Arial", 14), bg="navy",
                                 fg="white")
        self.submit4.grid(row=15, column=0, padx=10, pady=10, sticky=tk.N, ipadx=50)

        self.submit5 = tk.Button(self, text="Create Resume", command=self.resume_create, font=("Arial", 14), bg="navy",
                                 fg="white")
        self.submit5.grid(row=16, column=0, padx=10, pady=10, sticky=tk.N, ipadx=50)
        self.update()
        self.after(0, self.update_geometry)

    def submit_choice4(self):
        compN = self.comp.get("1.0",'end-1c')
        jobN = self.job.get("1.0",'end-1c')
        locationN = self.location.get("1.0",'end-1c')
        datesN = self.dates.get("1.0",'end-1c')
        compDescN = self.compDesc.get("1.0",'end-1c')
        jobDescN = self.jobDesc.get("1.0",'end-1c')
        achievementsN = self.achievements.get("1.0",'end-1c')
        self.experiences.append([compN, jobN, locationN, datesN, compDescN, jobDescN, achievementsN])
        self.comp.delete("1.0","end")
        self.job.delete("1.0","end")
        self.location.delete("1.0","end")
        self.dates.delete("1.0","end")
        self.compDesc.delete("1.0","end")
        self.jobDesc.delete("1.0","end")
        self.achievements.delete("1.0","end")

    def clear_all(self):
        for widget in self.winfo_children():
            widget.destroy()

    def resume_create(self):
        self.clear_all()
        print(self.resumeLevel)

        if self.resumeLevel == "Entry Level":
            self.cv = Resume.entryLevel(self.name, self.city, self.state, self.zip_code, self.phone, self.email,
                                        self.job, self.skills, self.education, self.experiences)
        elif self.resumeLevel == "Mid Level":
            self.cv = Resume.midLevel(self.name, self.city, self.state, self.zip_code, self.phone, self.email, self.job,
                                      self.skills, self.education, self.experiences)
        elif self.resumeLevel == "Senior Level":
            self.cv = Resume.seniorLevel(self.name, self.city, self.state, self.zip_code, self.phone, self.email,
                                         self.job, self.skills, self.education, self.experiences)
        elif self.resumeLevel == "Executive Level":
            self.cv = Resume.executiveLevel(self.name, self.city, self.state, self.zip_code, self.phone, self.email,
                                            self.job, self.skills, self.education, self.experiences)
        # create label for adjustments and text box for user input
        self.option = tk.Label(self,
                               text="Generated Resume",
                               font=("Arial", 14), fg="navy")
        self.option.grid(row=0, column=0, padx=10, pady=10, sticky=tk.W)
        # create scrolled text widget to display the resume
        self.resume = tkst.ScrolledText(self, wrap="word", font=("Arial", 10), bg="#F9F8F8")
        self.resume.insert(tk.END, self.cv)
        self.resume.grid(row=1, column=0, padx=10, pady=10, sticky=tk.NSEW)

        # create label for adjustments and text box for user input
        self.option = tk.Label(self,
                               text="Write any adjustments you would like in the box below, if none then just press save.",
                               font=("Arial", 14), fg="navy")
        self.option.grid(row=2, column=0, padx=10, pady=10, sticky=tk.W)

        self.adjustments = tk.Text(self, height=3, width=80, font=("Arial", 14), bg="#F9F8F8", wrap="word")
        self.adjustments.grid(row=3, column=0, padx=10, pady=5)

        # create scrollbar for the adjustments text widget
        scrollbar = tk.Scrollbar(self, command=self.adjustments.yview)
        scrollbar.grid(row=3, column=1, sticky='nsew')
        self.adjustments['yscrollcommand'] = scrollbar.set

        # create button to rewrite the resume based on user adjustments
        self.adjust = tk.Button(self, text="Rewrite", command=self.rewrite_resume, font=("Arial", 14), bg="navy",
                                 fg="white")
        self.adjust.grid(row=4, column=0, padx=10, pady=10, sticky=tk.W, ipadx=50)

        # create button to save the resume to a file
        self.save = tk.Button(self, text="Save", command=self.save_resume, font=("Arial", 14), bg="navy",
                                 fg="white")
        self.save.grid(row=4, column=0, padx=10, pady=10, sticky=tk.E, ipadx=50)

        # update geometry and display the form
        self.update()
        self.after(0, self.update_geometry)

    def new_resume_show(self, res):
        self.clear_all()

        # create label for adjustments and text box for user input
        self.option = tk.Label(self,
                               text="Generated Resume",
                               font=("Arial", 14), fg="navy")
        self.option.grid(row=0, column=0, padx=10, pady=10, sticky=tk.W)
        # create scrolled text widget to display the resume
        self.resume = tkst.ScrolledText(self, wrap="word", font=("Arial", 14), bg="#F9F8F8")
        self.resume.insert('1.0', res)
        self.resume.grid(row=1, column=0, padx=10, pady=10, sticky=tk.NSEW)

        # create label for adjustments and text box for user input
        self.option = tk.Label(self,
                               text="Write any adjustments you would like in the box below, if none then just press save.",
                               font=("Arial", 14), fg="navy")
        self.option.grid(row=2, column=0, padx=10, pady=10, sticky=tk.W)

        self.adjustments = tk.Text(self, height=3, width=80, font=("Arial", 14), bg="#F9F8F8", wrap="word")
        self.adjustments.grid(row=3, column=0, padx=10, pady=5)

        # create scrollbar for the adjustments text widget
        scrollbar = tk.Scrollbar(self, command=self.adjustments.yview)
        scrollbar.grid(row=3, column=1, sticky='nsew')
        self.adjustments['yscrollcommand'] = scrollbar.set

        # create button to rewrite the resume based on user adjustments
        self.adjust = tk.Button(self, text="Rewrite", command=self.rewrite_resume, font=("Arial", 14), bg="navy",
                                fg="white")
        self.adjust.grid(row=4, column=0, padx=10, pady=10, sticky=tk.W, ipadx=50)

        # create button to save the resume to a file
        self.save = tk.Button(self, text="Save", command=self.save_resume, font=("Arial", 14), bg="navy",
                              fg="white")
        self.save.grid(row=4, column=0, padx=10, pady=10, sticky=tk.E, ipadx=50)

        # update geometry and display the form
        self.update()
        self.after(0, self.update_geometry)

    def rewrite_resume(self):
        self.toAdjust = self.adjustments.get("1.0",'end-1c')
        self.adjustedResume = Resume.adjust(self.toAdjust, self.cv)
        self.new_resume_show(self.adjustedResume)

    def save_resume(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".doc", filetypes=[("Word Document", "*.docx")])
        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.size = Pt(16)
        if file_path:
            # Save the file
            with open(file_path, "a+") as file:
                if self.adjustedResume != '':
                    document.add_paragraph(self.adjustedResume)
                    document.save(file_path)
                else:
                    document.add_paragraph(self.cv)
                    document.save(file_path)

            # Notify the user that the file has been saved
            messagebox.showinfo("Resume saved", f"The resume has been saved at {file_path}")
        self.clear_all()
        self.experiences = []
        self.create_widgets()

root = tk.Tk()
app = Application(master=root)
app.mainloop()